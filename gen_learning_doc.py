# -*- coding: utf-8 -*-
# יצירת מסמך וורד — סיכום למידה

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# RTL for the whole document
for section in doc.sections:
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

style = doc.styles['Normal']
style.font.size = Pt(11)
style.font.name = 'David'
style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def add_title(text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in h.runs:
        run.font.name = 'David'

def add_p(text, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    run.font.name = 'David'
    run.font.size = Pt(11)
    run.bold = bold
    return p

def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    run.font.name = 'David'
    run.font.size = Pt(11)

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    # headers
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    # rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()  # spacing

# ==========================================
# תוכן המסמך
# ==========================================

add_title('סיכום למידה — המרת תוכניות לכתב כמויות', 1)
add_p('מסמך זה מסכם את כל מה שלמדנו מהצלבה עם פרויקט בת ים (בית ספר 24 כיתות)')

# --- חלק 1 ---
add_title('חלק 1: מה עשינו', 2)
add_p('לקחנו פרויקט אמיתי (בית ספר בת ים) עם תוכניות מלאות + כתב כמויות אמיתי (220 עמודים).')
add_p('קראנו 10 סוגי תוכניות: חשמל, אדריכלות, אינסטלציה, מיזוג, חתכים, תקרות, קונסטרוקציה, איטום, תקשורת, בקרה.')
add_p('ניסינו לייצר כתב כמויות לבד, הצלבנו מול האמיתי, ובנינו 7 כללים.')
add_p('גם בנינו סקריפט שמחלץ נתונים אוטומטית מ-PDF של תוכנית.')

# --- חלק 2 ---
add_title('חלק 2: 7 כללים', 2)

rules = [
    ('כלל 1: המקרא הוא המפתח',
     'רק מה שמופיע במקרא (Legend) של התוכנית נכנס לכתב הכמויות. אם זה לא במקרא — לא נכנס.'),
    ('כלל 2: כל תוכנית = נושא אחד',
     'תוכנית בינוי = מחיצות ודלתות בלבד. תוכנית חשמל = חשמל ותקשורת בלבד. לא לערבב!'),
    ('כלל 3: שמות סעיפים — להעתיק בדיוק מהמקרא',
     'לא להוסיף פרטים שלא כתובים. שינוי בשם = חיפוש סמל לא נכון = ספירה שגויה.'),
    ('כלל 4: קירות = מ"ר (לא מ"א!)',
     'כתב כמויות מודד קירות במ"ר = אורך קיר x גובה קומה. גובה קומה מופיע רק בחתך (PS).'),
    ('כלל 5: כמויות — לספור/למדוד מהתוכנית',
     'לא לנחש. בחשמל: הכמויות בטבלאות המעגלים. באדריכלות: למדוד אורכי קירות ולספור דלתות.'),
    ('כלל 6: להפריד כלל מנתון',
     '"קירות = מ"ר" — כלל קבוע. "גובה קומה = 280" — נתון ספציפי. בכל פרויקט חדש: לשאול, לא להניח.'),
    ('כלל 7: ביקורת עצמית אחרי כל למידה',
     'אחרי כל הצלבה — לחזור לכתבי כמויות קודמים ולבדוק שגיאות. הלמידה מצטברת.'),
]

for title, desc in rules:
    add_p(title, bold=True)
    add_p(desc)

# --- חלק 3 ---
add_title('חלק 3: איך ניגשים לכל סוג תוכנית', 2)

plan_types = [
    ('תוכנית בינוי (מחיצות)', [
        'מוצאים את המקרא — סוגי מחיצות ודלתות',
        'מודדים אורכי מחיצות לפי קנה מידה',
        'סופרים דלתות',
        'חייבים גובה קומה מחתך (PS) כדי להמיר מ"א למ"ר',
        'פרק 04 (קירות), פרק 06 (דלתות), פרק 14 (גבס)',
    ]),
    ('תוכנית חשמל', [
        'מוצאים 3 מקראות: חשמל, מאור, צבעי תעלות',
        'טבלאות מעגלים = מקור הכמויות — לא צריך לספור סמלים',
        'מודדים תעלות לפי קנה מידה, סופרים ארונות ולוחות',
        'פרק 08 (חשמל), פרק 18 (תקשורת), פרק 35 (בקרה)',
    ]),
    ('תוכנית אינסטלציה', [
        'מקרא צבעי צנרת: ירוק=ביוב, אדום=מים חמים, כחול=מים קרים',
        'מודדים אורכי צנרת, סופרים כלים סניטריים',
        'פרק 07',
    ]),
    ('תוכנית מיזוג אוויר', [
        'טבלת מעגלים = מקור הכמויות (BTU ליחידה)',
        'סופרים יחידות פנימיות וחיצוניות',
        'פרק 15',
    ]),
    ('חתכים (PS)', [
        'קוראים גובה כל קומה — חובה לפני חישוב מ"ר קירות',
        'מזהים חללים כפולים ועובי תקרות',
    ]),
    ('תקרות (PC)', [
        'מזהים סוגי תקרות אקוסטיות, מודדים שטח כל סוג',
        'פרק 14',
    ]),
    ('קונסטרוקציה', [
        'הכל מ"ק — חוץ מפלדה (טון) ותקרות (מ"ר)',
        'מפרידים לפי עובי (קיר 20 סמ = סעיף אחד, 25 סמ = סעיף אחר)',
        'פרק 02',
    ]),
    ('איטום', [
        'הכל מ"ר — שטחים לאטום',
        'מפרידים לפי סוג ומיקום (גג, יסודות, מקלחות)',
        'פרק 05',
    ]),
    ('תקשורת + בקרה', [
        'תקשורת (פרק 18) = המשך פרק 08. כבלי CAT 7A, ארונות rack',
        'בקרה (פרק 35) = גילוי אש/פריצה, כריזה. מופיע בתוכניות חשמל',
    ]),
]

for title, steps in plan_types:
    add_p(title, bold=True)
    for step in steps:
        add_bullet(step)

# --- חלק 4 ---
add_title('חלק 4: טעויות שעשינו ולמדנו מהן', 2)

add_table(
    ['מס', 'טעות', 'מה למדנו'],
    [
        ['1', 'הכנסנו פריטים שלא במקרא (גבס, ריצוף, חשמל בתוכנית בינוי)', 'כלל 1+2: רק מה שבמקרא, רק של הנושא'],
        ['2', 'כתבנו "9 יחידות מחיצה" במקום למדוד מ"א', 'כלל 4+5: קירות=מ"ר, למדוד לא לנחש'],
        ['3', 'הוספנו "H=40" שלא כתוב במקרא', 'כלל 3: להעתיק בדיוק'],
        ['4', 'ספרנו 3 לוחות חשמל במקום 4', 'להצליב ספירה ויזואלית עם טקסט'],
    ]
)

# --- חלק 5 ---
add_title('חלק 5: שאלות לכל פרויקט חדש', 2)
checklist = [
    'מה גובה כל קומה? — לבדוק בחתך (PS)',
    'יש חללים כפולים? — לבדוק בחתך (PS)',
    'מה סוגי המחיצות ואיך מבדילים ביניהם? — מקרא + צבע/עובי',
    'שלב א או שלב ב? — כתב הכמויות מפריד',
    'איזה מחירון? שנה?',
    'מה כל סימון (PC/PL/PS/PE) אומר? — לבדוק כותרת, לא להניח',
]
for item in checklist:
    add_bullet(item)

# --- חלק 6 ---
add_title('חלק 6: מה הלאה', 2)
add_p('1. לירן בודק ומתקן את כתב הכמויות של הסוללים')
add_p('2. אנחנו מתקנים לפי ההערות שלו')
add_p('3. בונים מנוע אוטומטי: PDF נכנס => כתב כמויות יוצא')

# שמירה
output = r'C:\Users\LENOVO\Desktop\cluade\contractor-demo\ללירן\סיכום-למידה.docx'
doc.save(output)
print(f'Saved: {output}')
